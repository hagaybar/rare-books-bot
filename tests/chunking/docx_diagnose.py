from docx import Document
from docx.oxml.ns import qn
import pathlib
import hashlib


def diagnose_docx_images(path: str | pathlib.Path):
    """Diagnose how images are structured in a DOCX file."""
    document = Document(path)

    print(f"\n=== DOCX Image Structure Analysis for {path} ===\n")

    # Collect all image relationships
    image_rels = {}
    for rel_id, rel in document.part.rels.items():
        if "image" in rel.reltype:
            image_hash = hashlib.md5(rel.target_part.blob).hexdigest()[:8]
            image_rels[rel_id] = {
                'type': rel.target_part.content_type,
                'size': len(rel.target_part.blob),
                'hash': image_hash,
            }

    print(f"Total image relationships found: {len(image_rels)}")
    for rid, info in image_rels.items():
        print(f"  {rid}: {info['type']}, {info['size']} bytes, hash: {info['hash']}")

    print(f"\nAnalyzing {len(document.paragraphs)} paragraphs for images:\n")

    # Check each paragraph
    images_found_total = 0
    paragraph_image_map = {}  # Track which paragraphs have which images

    for idx, paragraph in enumerate(document.paragraphs, start=1):
        # Search for blips using the simpler xpath without namespaces
        blips = paragraph._element.xpath(".//a:blip")
        drawings = paragraph._element.xpath(".//w:drawing")
        picts = paragraph._element.xpath(".//w:pict")

        if blips or drawings or picts:
            print(f"Paragraph {idx}:")
            if paragraph.text.strip():
                print(f"  Text preview: {paragraph.text[:50]}...")
            else:
                print(f"  (No text - image only paragraph)")

            if drawings:
                print(f"  Found {len(drawings)} w:drawing element(s)")
            if picts:
                print(f"  Found {len(picts)} w:pict element(s)")

            # For blips, show the rIds
            if blips:
                print(f"  Found {len(blips)} a:blip element(s)")
                paragraph_rids = []
                for blip in blips:
                    rId = blip.get(qn("r:embed"))
                    if rId:
                        paragraph_rids.append(rId)
                        if rId in image_rels:
                            info = image_rels[rId]
                            print(f"    → {rId}: hash={info['hash']}, size={info['size']}")
                            images_found_total += 1
                        else:
                            print(f"    → {rId}: (not in image relationships)")

                paragraph_image_map[idx] = paragraph_rids

            # Also check the raw XML for debugging
            xml_snippet = paragraph._element.xml[:200]
            if 'drawing' in xml_snippet or 'pict' in xml_snippet:
                print(f"  XML snippet: {xml_snippet}...")

    print(f"\nSummary: Found {images_found_total} image references across all paragraphs")

    # Check document body for all images
    body_element = document.element.body
    all_body_blips = body_element.xpath(".//a:blip")
    print(f"\nTotal images in document body: {len(all_body_blips)}")

    # Find orphaned images (in body but not in any paragraph)
    all_found_rids = set()
    for rids in paragraph_image_map.values():
        all_found_rids.update(rids)

    print("\nChecking for images not associated with paragraphs:")
    for blip in all_body_blips:
        rId = blip.get(qn("r:embed"))
        if rId and rId not in all_found_rids:
            print(f"  Found orphaned image: {rId}")
            # Try to find its parent structure
            parent = blip.getparent()
            while parent is not None and parent.tag not in [qn('w:p'), qn('w:body')]:
                parent = parent.getparent()
            if parent is not None and parent.tag == qn('w:p'):
                # Find which paragraph this is
                for idx, para in enumerate(document.paragraphs, start=1):
                    if para._element == parent:
                        print(f"    → Actually belongs to paragraph {idx}")
                        break

    # Show the final mapping
    print("\nFinal paragraph-to-image mapping:")
    for para_idx, rids in sorted(paragraph_image_map.items()):
        print(f"  Paragraph {para_idx}: {', '.join(rids)}")


# Usage:
# diagnose_docx_images("path/to/your/docx_new_example.docx")
# Usage:
diagnose_docx_images(
    "/home/hagaybar/projects/Multi-Source_RAG_Platform/data/projects/test_minimal_data/input/raw/docx/docx_new_example.docx"
)
