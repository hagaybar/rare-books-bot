from __future__ import annotations

import pathlib
from pathlib import Path
from typing import List, Tuple, Dict
from docx import Document
from docx.oxml.ns import qn
import logging
import hashlib

from scripts.utils.image_utils import (
    infer_project_root,
    get_project_image_dir,
    save_image_blob,
    generate_image_filename,
)
from scripts.utils.logger import LoggerManager

# Initialize logger with proper file output  
logger = LoggerManager.get_logger("docx_ingestor")


def load_docx(path: str | pathlib.Path, run_id: str | None = None) -> List[Tuple[str, dict]]:
    """Extract text and image references from a .docx file as
    (text, metadata) chunks."""
    if not isinstance(path, Path):
        path = Path(path)

    document = Document(path)
    project_root = infer_project_root(path)
    image_dir = get_project_image_dir(project_root.name)
    doc_id = str(path)
    
    # Initialize logger with run_id support
    docx_logger = LoggerManager.get_logger("docx_loader", run_id=run_id)
    
    docx_logger.debug("Writing image to directory", extra={"run_id": run_id, "doc_id": doc_id, "image_dir": str(image_dir)} if run_id else {"doc_id": doc_id, "image_dir": str(image_dir)})

    segments: List[Tuple[str, dict]] = []

    # Track saved images to avoid duplicates
    saved_images: Dict[str, str] = {}  # rId -> saved_path
    img_counter = 0  # Global counter for image numbering

    for para_idx, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        meta = {
            "doc_type": "docx",
            "paragraph_number": para_idx,
            "source_filepath": str(path),
            "doc_id": doc_id,
        }

        image_paths = []

        # Look for blips in this paragraph
        blips = paragraph._element.xpath(".//a:blip")

        for blip in blips:
            rId = blip.get(qn("r:embed"))
            if not rId:
                continue

            # Check if we already saved this image
            if rId in saved_images:
                # Reuse the existing path
                image_paths.append(saved_images[rId])
                docx_logger.debug(
                    "Reusing image",
                    extra={"run_id": run_id, "doc_id": doc_id, "paragraph_number": para_idx, "image_rid": rId, "image_path": saved_images[rId]} if run_id else {"doc_id": doc_id, "paragraph_number": para_idx, "image_rid": rId, "image_path": saved_images[rId]}
                )
            else:
                # Get the image part
                image_part = document.part.related_parts.get(rId)
                if not image_part:
                    continue

                # Save the new image
                img_name = generate_image_filename(
                    doc_id=doc_id,
                    page_number=para_idx,
                    img_index=img_counter,
                )
                img_path = image_dir / img_name
                save_image_blob(image_part.blob, img_path)

                try:
                    rel_path = img_path.resolve().relative_to(
                        (project_root / "input").resolve()
                    )
                    saved_path = str(rel_path)
                except ValueError:
                    saved_path = f"cache/images/{img_path.name}"

                saved_images[rId] = saved_path
                image_paths.append(saved_path)
                img_counter += 1

                docx_logger.info(
                    "Saved new image",
                    extra={"run_id": run_id, "doc_id": doc_id, "paragraph_number": para_idx, "image_rid": rId, "image_path": saved_path} if run_id else {"doc_id": doc_id, "paragraph_number": para_idx, "image_rid": rId, "image_path": saved_path}
                )

        # Add image paths to metadata if any were found
        if image_paths:
            meta["image_paths"] = image_paths
            docx_logger.debug(
                "Extracted images from paragraph",
                extra={"run_id": run_id, "doc_id": doc_id, "paragraph_number": para_idx, "image_count": len(image_paths), "image_paths": image_paths} if run_id else {"doc_id": doc_id, "paragraph_number": para_idx, "image_count": len(image_paths), "image_paths": image_paths}
            )

        # Create segment if there's text or images
        if text or image_paths:
            segments.append((text or "[Image-only content]", meta))
        else:
            docx_logger.debug(
                "Paragraph skipped - no content",
                extra={"run_id": run_id, "doc_id": doc_id, "paragraph_number": para_idx} if run_id else {"doc_id": doc_id, "paragraph_number": para_idx}
            )

    # Debug: Log all segments with images
    image_segments = [(idx, meta) for idx, (text, meta) in enumerate(segments) if 'image_paths' in meta]
    if image_segments:
        docx_logger.debug(
            "All segments with images summary",
            extra={"run_id": run_id, "doc_id": doc_id, "image_segment_count": len(image_segments), "image_segments": [f"Segment {idx}: Paragraph {meta['paragraph_number']}, Images: {meta['image_paths']}" for idx, meta in image_segments]} if run_id else {"doc_id": doc_id, "image_segment_count": len(image_segments), "image_segments": [f"Segment {idx}: Paragraph {meta['paragraph_number']}, Images: {meta['image_paths']}" for idx, meta in image_segments]}
        )

    image_attached_chunks = sum('image_paths' in m for _, m in segments)
    docx_logger.info(
        "Extraction complete",
        extra={"run_id": run_id, "doc_id": doc_id, "total_segments": len(segments), "image_attached_chunks": image_attached_chunks, "file_name": path.name} if run_id else {"doc_id": doc_id, "total_segments": len(segments), "image_attached_chunks": image_attached_chunks, "file_name": path.name}
    )

    # Add tables
    for tbl_idx, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                rows.append(" | ".join(row_cells))
        if rows:
            tbl_text = "\n".join(rows)
            meta = {
                "doc_type": "docx",
                "table_number": tbl_idx,
                "source_filepath": str(path),
                "doc_id": doc_id,
            }
            segments.append((tbl_text, meta))

    # Final verification
    final_image_segments = [(i, meta) for i, (text, meta) in enumerate(segments) if 'image_paths' in meta]
    docx_logger.debug(
        "Final segment verification",
        extra={"run_id": run_id, "doc_id": doc_id, "total_segments": len(segments), "final_image_segments": [f"Segment {i}: {meta.get('doc_type')} para {meta.get('paragraph_number', 'N/A')}, images: {meta['image_paths']}" for i, meta in final_image_segments]} if run_id else {"doc_id": doc_id, "total_segments": len(segments), "final_image_segments": [f"Segment {i}: {meta.get('doc_type')} para {meta.get('paragraph_number', 'N/A')}, images: {meta['image_paths']}" for i, meta in final_image_segments]}
    )

    return segments
