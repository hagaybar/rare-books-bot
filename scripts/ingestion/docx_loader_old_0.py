from __future__ import annotations

import pathlib
from pathlib import Path
from typing import List, Tuple
from docx import Document
from docx.oxml.ns import qn
import logging
import os

from scripts.utils.image_utils import (
    infer_project_root,
    get_project_image_dir,
    save_image_blob,
    generate_image_filename,
)

logger = logging.getLogger("docx_ingestor")


def load_docx(path: str | pathlib.Path) -> List[Tuple[str, dict]]:
    """Extract text and image references from a .docx file as (text, metadata)
    chunks."""
    if not isinstance(path, Path):
        path = Path(path)

    document = Document(path)
    project_root = infer_project_root(path)
    image_dir = get_project_image_dir(project_root.name)
    doc_id = str(path)

    print(f"[loader] Writing image to: {image_dir}")

    segments: List[Tuple[str, dict]] = []

    for para_idx, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        meta = {
            "doc_type": "docx",
            "paragraph_number": para_idx,
            "source_filepath": str(path),
            "doc_id": doc_id,
        }

        image_paths = []
        img_count = 0

        for run in paragraph.runs:
            blips = run._element.xpath(".//a:blip")
            for blip in blips:
                rId = blip.get(qn("r:embed"))
                image_part = document.part.related_parts.get(rId)
                if not image_part:
                    continue

                img_name = generate_image_filename(
                    doc_id=doc_id,
                    page_number=para_idx,
                    img_index=img_count,
                )
                img_path = image_dir / img_name
                save_image_blob(image_part.blob, img_path)
                if image_paths:
                    print(
                        f"[DEBUG] Paragraph {para_idx} → extracted "
                        f"{len(image_paths)} image(s): {image_paths}"
                    )
                    logger.info(
                        f"[DOCX] [DEBUG] Paragraph {para_idx} → extracted "
                        f"{len(image_paths)} image(s): {image_paths}"
                    )

                try:
                    rel_path = img_path.resolve().relative_to(
                        (project_root / "input").resolve()
                    )
                    image_paths.append(str(rel_path))
                except ValueError:
                    image_paths.append(f"cache/images/{img_path.name}")

                img_count += 1

        if image_paths:
            meta["image_paths"] = image_paths

        if not (text or image_paths):
            print(
                f"[DEBUG] Paragraph {para_idx} was skipped — no text and "
                f"no images recorded."
            )
            logger.debug(
                f"[DOCX] [DEBUG] Paragraph {para_idx} was skipped — no text and "
                f"no images recorded."
            )

        # ✅ Fallback: if no image_paths were recorded but images were saved
        # for this paragraph
        if not image_paths:
            # Matches filenames like: docx_new_example_page15_img0.png
            pattern = f"{path.stem}_page{para_idx}_img*.png"
            potential_images = sorted(image_dir.glob(pattern))

            if potential_images:
                image_paths = []
                for p in potential_images:
                    try:
                        rel = p.resolve().relative_to(
                            (project_root / "input").resolve()
                        )
                        image_paths.append(str(rel))
                    except ValueError:
                        # Fallback for strange paths
                        image_paths.append(f"cache/images/{p.name}")
                meta["image_paths"] = image_paths

        if text or image_paths:
            segments.append((text or "[Image-only content]", meta))

    print(
        f"[INFO] Extracted {sum('image_paths' in m for _, m in segments)} "
        f"image-attached chunks from {path.name}"
    )

    # Add tables
    for tbl_idx, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            row_cells = [
                cell.text.strip() for cell in row.cells if cell.text.strip()
            ]
            if row_cells:
                rows.append(" | ".join(row_cells))
        if rows:
            tbl_text = "\n".join(rows)
            meta = {
                "doc_type": "docx",
                "table_number": tbl_idx,
                "source_filepath": str(path),
                "doc_id": doc_id,
            }
            segments.append((tbl_text, meta))

    return segments
